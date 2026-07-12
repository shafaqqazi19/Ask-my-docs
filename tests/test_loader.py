from src.ingestion.loader import load_documents


def test_docx_loader_extracts_paragraphs_and_tables(tmp_path):
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_heading("Vacation Policy", level=1)
    doc.add_paragraph("Employees accrue 1.5 days of PTO per month worked.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Tenure"
    table.rows[0].cells[1].text = "Max Carryover"
    table.rows[1].cells[0].text = "0-2 years"
    table.rows[1].cells[1].text = "5 days"
    doc.save(tmp_path / "policy.docx")

    docs = load_documents(tmp_path)
    assert len(docs) == 1
    assert "1.5 days of PTO" in docs[0].text
    assert "Max Carryover" in docs[0].text  # table content preserved


def test_html_loader_strips_scripts_and_styles(tmp_path):
    html = """
    <html><head><style>.x{color:red}</style><script>alert('x')</script></head>
    <body><h1>Shipping FAQ</h1><p>Standard shipping takes 3-5 business days.</p></body>
    </html>
    """
    (tmp_path / "page.html").write_text(html, encoding="utf-8")

    docs = load_documents(tmp_path)
    assert len(docs) == 1
    assert "Shipping FAQ" in docs[0].text
    assert "3-5 business days" in docs[0].text
    assert "alert" not in docs[0].text
    assert "color:red" not in docs[0].text


def test_csv_loader_serializes_rows_as_readable_text(tmp_path):
    csv_content = "sku,name,price_usd\nA100,Widget Pro,19.99\nA101,Widget Mini,9.99\n"
    (tmp_path / "inventory.csv").write_text(csv_content, encoding="utf-8")

    docs = load_documents(tmp_path)
    assert len(docs) == 1
    assert "sku: A100" in docs[0].text
    assert "name: Widget Pro" in docs[0].text
    assert "price_usd: 19.99" in docs[0].text


def test_corrupt_file_is_skipped_not_fatal(tmp_path):
    (tmp_path / "broken.docx").write_text("not a real docx", encoding="utf-8")
    (tmp_path / "good.txt").write_text("This one is fine.", encoding="utf-8")

    docs = load_documents(tmp_path)
    doc_ids = [d.doc_id for d in docs]
    assert "good.txt" in doc_ids
    assert "broken.docx" not in doc_ids


def test_unsupported_extension_is_silently_skipped(tmp_path):
    (tmp_path / "image.png").write_bytes(b"\x89PNG\r\n")
    (tmp_path / "good.txt").write_text("readable", encoding="utf-8")

    docs = load_documents(tmp_path)
    assert len(docs) == 1
    assert docs[0].doc_id == "good.txt"
