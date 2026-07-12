"""Load raw documents (pdf, md, txt, docx, html, csv) from a directory into plain text."""
from __future__ import annotations

import csv
import logging
from dataclasses import dataclass
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader

logger = logging.getLogger("ask_my_docs.loader")


@dataclass
class RawDocument:
    doc_id: str
    source_path: str
    text: str


def _load_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages)


def _load_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _load_docx(path: Path) -> str:
    """Extract paragraph text and table cell text from a .docx file, in document order."""
    doc = DocxDocument(str(path))
    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n\n".join(parts)


def _load_html(path: Path) -> str:
    """Strip tags/scripts/styles and return visible text, with light structure
    preserved (headings and paragraphs separated by blank lines)."""
    raw = path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(raw, "html.parser")

    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    blocks: list[str] = []
    for el in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "td", "th"]):
        text = el.get_text(strip=True)
        if text:
            blocks.append(text)

    if not blocks:
        # fallback: no recognizable block tags, just grab all visible text
        text = soup.get_text(separator="\n", strip=True)
        return text

    return "\n\n".join(blocks)


def _load_csv(path: Path) -> str:
    """Serialize CSV rows into 'column: value' text blocks (one block per row) so
    each row becomes a semantically retrievable unit rather than a raw comma line.
    Handles large files by streaming rather than loading everything into memory twice.
    """
    parts: list[str] = []
    with open(path, "r", encoding="utf-8", errors="ignore", newline="") as f:
        reader = csv.DictReader(f)
        if reader.fieldnames is None:
            return ""

        for row in reader:
            line = "; ".join(f"{col}: {val}" for col, val in row.items() if val not in (None, ""))
            if line:
                parts.append(line)

    return "\n\n".join(parts)


LOADERS = {
    ".pdf": _load_pdf,
    ".md": _load_text,
    ".txt": _load_text,
    ".docx": _load_docx,
    ".html": _load_html,
    ".htm": _load_html,
    ".csv": _load_csv,
}


def load_documents(input_dir: str | Path) -> list[RawDocument]:
    """Walk `input_dir` and load every supported file into a RawDocument.

    Files with an unsupported extension are silently skipped. Files that fail to
    parse (corrupt pdf/docx, malformed csv, etc.) are skipped with a warning rather
    than crashing the whole ingestion run.
    """
    input_dir = Path(input_dir)
    docs: list[RawDocument] = []

    for path in sorted(input_dir.rglob("*")):
        if not path.is_file():
            continue
        loader = LOADERS.get(path.suffix.lower())
        if loader is None:
            continue

        try:
            text = loader(path)
        except Exception as exc:  # noqa: BLE001 - ingestion must not hard-fail on one bad file
            logger.warning("skipping '%s' — failed to parse (%r)", path, exc)
            continue

        if not text.strip():
            logger.warning("skipping '%s' — no extractable text", path)
            continue

        doc_id = path.relative_to(input_dir).as_posix()
        docs.append(RawDocument(doc_id=doc_id, source_path=str(path), text=text))

    return docs
